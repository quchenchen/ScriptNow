"""Unit tests for :mod:`app.services.document_parser`."""
from __future__ import annotations

from pathlib import Path

import pytest

from app.services.document_parser import (
    SUPPORTED_EXTENSIONS,
    UnsupportedDocument,
    is_supported,
    parse,
)


def test_supported_extensions_include_common_formats():
    assert ".docx" in SUPPORTED_EXTENSIONS
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".txt" in SUPPORTED_EXTENSIONS
    assert ".md" in SUPPORTED_EXTENSIONS


def test_is_supported_case_insensitive():
    assert is_supported("hello.DOCX")
    assert is_supported("plot.MD")
    assert not is_supported("weird.rtf")


def test_parse_txt(tmp_path: Path):
    p = tmp_path / "seed.txt"
    p.write_text("这是文本文档的内容。\n第二段。", encoding="utf-8")
    text = parse(p)
    assert "文本文档" in text
    assert "第二段" in text


def test_parse_md(tmp_path: Path):
    p = tmp_path / "outline.md"
    p.write_text("# 标题\n\n正文行 1", encoding="utf-8")
    text = parse(p)
    assert "标题" in text


def test_parse_gbk_txt_falls_back(tmp_path: Path):
    p = tmp_path / "legacy.txt"
    p.write_bytes("旧文档内容".encode("gbk"))
    text = parse(p)
    assert "旧文档" in text


def test_parse_docx(tmp_path: Path):
    from docx import Document
    p = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("段落一：介绍主人公。")
    doc.add_paragraph("段落二：冲突。")
    doc.save(str(p))
    text = parse(p)
    assert "段落一" in text
    assert "段落二" in text


def test_parse_unsupported_extension_raises(tmp_path: Path):
    p = tmp_path / "weird.xyz"
    p.write_text("nope")
    with pytest.raises(UnsupportedDocument):
        parse(p)
