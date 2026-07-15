"""Extract plain text from uploaded documents.

Supports ``.docx / .pdf / .txt / .md``. Chosen for coverage of the everyday
formats writers use; ``.doc`` (legacy Word) requires system tools like
antiword and is out of scope. Extraction is per-format:

- ``.docx`` — python-docx paragraph iteration
- ``.pdf`` — pypdf page iteration
- ``.txt`` / ``.md`` — UTF-8 read with fallback to GBK

All entry points return a single ``str`` blob. Callers should chunk with
:mod:`app.services.chunker` before persisting.
"""
from __future__ import annotations

from pathlib import Path


class UnsupportedDocument(ValueError):
    """Raised when the file extension is not supported."""


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # A lot of legacy Chinese ``.txt`` files are GBK
        return path.read_text(encoding="gbk", errors="ignore")


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            paragraphs.append(p.text)
    # Include tables too — some scripts have data in tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages: list[str] = []
    for p in reader.pages:
        try:
            pages.append(p.extract_text() or "")
        except Exception:  # pragma: no cover — corrupted page
            continue
    return "\n\n".join(pages)


_EXT_HANDLERS = {
    ".txt": _read_text_file,
    ".md": _read_text_file,
    ".docx": _read_docx,
    ".pdf": _read_pdf,
}

SUPPORTED_EXTENSIONS = tuple(_EXT_HANDLERS.keys())


def parse(path: str | Path) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    handler = _EXT_HANDLERS.get(ext)
    if handler is None:
        raise UnsupportedDocument(f"unsupported extension: {ext}")
    return handler(p)


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in _EXT_HANDLERS
