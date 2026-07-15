"""Export episodes as ``.docx`` in script-sheet layout.

Layout follows the format checker's contract:
- ``【场景N】location·time`` becomes a bold standalone heading paragraph
- ``△...`` action lines become italic paragraphs with left indent
- ``角色名：对白`` dialog lines become a bold speaker name + regular content

Uses python-docx (already in pyproject.toml).
"""
from __future__ import annotations

import io
import re
from collections.abc import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

_SCENE_HEAD = re.compile(r"^\s*【场景\s*(\d+)\s*】\s*(.*)$")
_ACTION = re.compile(r"^\s*△\s*(.*)$")
_DIALOG = re.compile(r"^\s*([^\s：:]{1,10})：(.+)$")


def _add_scene_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _add_action(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"△ {text}")
    run.italic = True
    run.font.size = Pt(11)


def _add_dialog(doc: Document, speaker: str, line: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{speaker}：")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(line).font.size = Pt(11)


def _add_plain(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.8)
    for r in p.runs:
        r.font.size = Pt(11)


def _iter_lines_as_paragraphs(doc: Document, lines: Iterable[str]) -> None:
    """Render one episode's lines into the doc in script-sheet layout."""
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        m = _SCENE_HEAD.match(line)
        if m:
            n = m.group(1)
            rest = (m.group(2) or "").strip()
            _add_scene_heading(doc, f"【场景{n}】{rest}")
            continue
        m = _ACTION.match(line)
        if m:
            _add_action(doc, m.group(1))
            continue
        m = _DIALOG.match(line)
        if m:
            _add_dialog(doc, m.group(1), m.group(2))
            continue
        _add_plain(doc, line)


def build_episode_text(scenes: list[dict]) -> str:
    """Assemble scene rows into the raw text a script-sheet expects.

    Kept public so callers (API layer) can share it with the format checker.
    """
    if not scenes:
        return ""
    chunks: list[str] = []
    for s in scenes:
        head = f"【场景{s['scene_number']}】{s.get('location', '') or ''}"
        if s.get("time"):
            head += f"·{s['time']}"
        body = s.get("content", "") or ""
        chunks.append(f"{head}\n{body}")
    return "\n\n".join(chunks)


def render_project(project: dict, episodes: list[dict]) -> bytes:
    """Render a whole project to a .docx byte string.

    ``project`` is the project row dict; ``episodes`` is a list of dicts with
    ``episode_number`` / ``title`` / ``scenes`` (list of scene rows).
    """
    doc = Document()

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(project.get("title") or "未命名剧本")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(
        f"{project.get('type', 'script')} · {len(episodes)} 集"
    )
    subtitle_run.font.size = Pt(11)

    doc.add_page_break()

    # Body
    for ep in episodes:
        ep_h = doc.add_paragraph()
        ep_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ep_h.paragraph_format.space_before = Pt(18)
        ep_run = ep_h.add_run(f"第{ep['episode_number']}集 {ep.get('title', '')}")
        ep_run.bold = True
        ep_run.font.size = Pt(16)

        text = build_episode_text(ep.get("scenes") or [])
        _iter_lines_as_paragraphs(doc, text.split("\n"))

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
