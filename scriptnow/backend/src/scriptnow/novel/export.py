import io
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from scriptnow.novel.contracts import NovelBlock


@dataclass(frozen=True, slots=True)
class NovelExportChapter:
    volume_title: str
    chapter_title: str
    blocks: tuple[NovelBlock, ...]


def render_novel_docx(*, project_name: str, chapters: tuple[NovelExportChapter, ...]) -> bytes:
    """Render Novel blocks without importing Script formatting or paragraph types."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    title = document.add_paragraph(project_name)
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_volume = None
    for index, chapter in enumerate(chapters):
        if index:
            document.add_page_break()
        if current_volume != chapter.volume_title:
            volume = document.add_paragraph(chapter.volume_title)
            volume.style = document.styles["Heading 1"]
            volume.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_volume = chapter.volume_title
        chapter_heading = document.add_paragraph(chapter.chapter_title)
        chapter_heading.style = document.styles["Heading 2"]
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for block in chapter.blocks:
            if block.type == "heading":
                paragraph = document.add_paragraph(block.text, style="Heading 3")
                paragraph.runs[0].font.name = "黑体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif block.type == "divider":
                paragraph = document.add_paragraph("＊ ＊ ＊")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].font.name = "宋体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            else:
                paragraph = document.add_paragraph(block.text)
                paragraph.paragraph_format.line_spacing = 1.5
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.runs[0].font.name = "宋体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if block.type == "prose":
                    paragraph.paragraph_format.first_line_indent = Cm(0.74)
                elif block.type == "quote":
                    paragraph.paragraph_format.left_indent = Cm(1.5)
                    paragraph.paragraph_format.right_indent = Cm(1.5)
                    paragraph.runs[0].italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def render_packaged_docx(
    *,
    project_name: str,
    synopsis: str,
    tags: tuple[str, ...],
    cover_image_bytes: bytes | None,
    chapters: tuple[NovelExportChapter, ...],
) -> bytes:
    """Render a DOCX with cover image, synopsis, tags, and manuscript."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    # Cover image
    if cover_image_bytes:
        image_stream = io.BytesIO(cover_image_bytes)
        cover_para = document.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_para.add_run()
        run.add_picture(image_stream, width=Cm(10))
        document.add_page_break()

    # Title
    title = document.add_paragraph(project_name)
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Synopsis
    if synopsis:
        synopsis_heading = document.add_paragraph("简介", style="Heading 1")
        synopsis_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        synopsis_text = document.add_paragraph(synopsis)
        synopsis_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tags
    if tags:
        document.add_paragraph()
        tag_para = document.add_paragraph(" · ".join(tags))
        tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

    # Manuscript
    current_volume = None
    for index, chapter in enumerate(chapters):
        if index:
            document.add_page_break()
        if current_volume != chapter.volume_title:
            volume = document.add_paragraph(chapter.volume_title)
            volume.style = document.styles["Heading 1"]
            volume.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_volume = chapter.volume_title
        chapter_heading = document.add_paragraph(chapter.chapter_title)
        chapter_heading.style = document.styles["Heading 2"]
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for block in chapter.blocks:
            if block.type == "heading":
                paragraph = document.add_paragraph(block.text, style="Heading 3")
                paragraph.runs[0].font.name = "黑体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif block.type == "divider":
                paragraph = document.add_paragraph("＊ ＊ ＊")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph = document.add_paragraph(block.text)
                paragraph.runs[0].font.name = "宋体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
