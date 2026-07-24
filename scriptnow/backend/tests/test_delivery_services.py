import hashlib
import io

import pytest
from docx import Document
from sqlalchemy import select

from scriptnow.novel.contracts import NovelBlock
from scriptnow.novel.delivery import NovelDeliveryError, NovelExportService
from scriptnow.novel.domain import NovelDocumentRevisionModel, NovelRevisionStatus
from scriptnow.novel.project import NovelStoryMapModel, initialize_novel_project
from scriptnow.platform.database import Database
from scriptnow.platform.models import ProjectMedium, ProjectModel, TenantModel
from scriptnow.platform.translation_contracts import TranslationUnit
from scriptnow.script.contracts import ScriptBlock
from scriptnow.script.delivery import ScriptDeliveryError, ScriptExportService
from scriptnow.script.domain import RevisionStatus, ScriptDocumentRevisionModel
from scriptnow.script.export import render_script_docx
from scriptnow.script.project import ScriptStoryMapModel, initialize_script_project


@pytest.fixture
async def delivery_data():
    database = Database.create("sqlite+aiosqlite:///:memory:")
    await database.create_schema()
    async with database.session() as session:
        tenant = TenantModel(name="Delivery Studio")
        session.add(tenant)
        await session.flush()
        script = ProjectModel(
            tenant_id=tenant.id,
            name="剧本项目",
            medium=ProjectMedium.SCRIPT,
            direction={"script_format": "chinese", "language": "zh-CN"},
        )
        novel = ProjectModel(
            tenant_id=tenant.id,
            name="小说项目",
            medium=ProjectMedium.NOVEL,
            direction={"language": "zh-CN"},
        )
        session.add_all([script, novel])
        await session.flush()
        await initialize_script_project(session, script)
        await initialize_novel_project(session, novel)
        script_map = (
            await session.scalars(
                select(ScriptStoryMapModel).where(ScriptStoryMapModel.project_id == script.id)
            )
        ).one()
        script_map.episodes = [
            {
                "id": "episode-1",
                "title": "第一集",
                "scenes": [
                    {"id": "scene-1", "title": "雨夜"},
                    {"id": "scene-2", "title": "清晨"},
                ],
            }
        ]
        novel_map = (
            await session.scalars(
                select(NovelStoryMapModel).where(NovelStoryMapModel.project_id == novel.id)
            )
        ).one()
        novel_map.volumes = [
            {
                "id": "volume-1",
                "title": "第一卷",
                "chapters": [
                    {"id": "chapter-1", "title": "来信"},
                    {"id": "chapter-2", "title": "回声"},
                ],
            }
        ]
        session.add(
            ScriptDocumentRevisionModel(
                project_id=script.id,
                scene_id="scene-1",
                revision_number=1,
                blocks=[
                    ScriptBlock(para_id="p1", type="slugline", text="内景 客厅 夜").model_dump(
                        mode="json"
                    ),
                    ScriptBlock(
                        para_id="p2", type="action", text="她打开信。充满疑惑。她沉默。"
                    ).model_dump(mode="json"),
                ],
                status=RevisionStatus.ADOPTED,
                idempotency_key="script-doc",
            )
        )
        session.add(
            NovelDocumentRevisionModel(
                project_id=novel.id,
                chapter_id="chapter-1",
                revision_number=1,
                blocks=[
                    NovelBlock(block_id="h1", type="heading", text="来信").model_dump(mode="json"),
                    NovelBlock(
                        block_id="b1", type="prose", text="雨落在窗沿，信纸发出轻响。"
                    ).model_dump(mode="json"),
                ],
                status=NovelRevisionStatus.ADOPTED,
                idempotency_key="novel-doc",
            )
        )
    yield database, tenant, script, novel
    await database.dispose()


@pytest.mark.asyncio
async def test_script_manifest_partial_scope_idempotency_and_failed_retry(delivery_data) -> None:
    database, tenant, script, _ = delivery_data
    calls = 0

    def flaky(**kwargs):
        nonlocal calls
        calls += 1
        if calls == 1:
            raise RuntimeError("temporary renderer failure")
        return render_script_docx(**kwargs)

    service = ScriptExportService(database, flaky)
    options = await service.options(tenant_id=tenant.id, project_id=script.id)
    assert options["episodes"][0]["selection"] == "partial"
    assert [item["selectable"] for item in options["episodes"][0]["scenes"]] == [True, False]
    failed = await service.export(
        tenant_id=tenant.id,
        project_id=script.id,
        scene_ids=("scene-1",),
        form="clean",
        idempotency_key="export-1",
    )
    assert failed.status == "failed" and failed.attempts == 1
    retried = await service.export(
        tenant_id=tenant.id,
        project_id=script.id,
        scene_ids=("scene-1",),
        form="clean",
        idempotency_key="export-1",
    )
    assert retried.id == failed.id and retried.status == "succeeded" and retried.attempts == 2
    assert retried.artifact_sha256 == hashlib.sha256(retried.artifact).hexdigest()
    replay = await service.export(
        tenant_id=tenant.id,
        project_id=script.id,
        scene_ids=("scene-1",),
        form="clean",
        idempotency_key="export-1",
    )
    assert replay.id == retried.id and replay.attempts == 2 and calls == 2
    with pytest.raises(ScriptDeliveryError, match="completed"):
        await service.export(
            tenant_id=tenant.id,
            project_id=script.id,
            scene_ids=("scene-2",),
            form="clean",
            idempotency_key="empty-scene",
        )


@pytest.mark.asyncio
async def test_novel_export_uses_independent_contract_and_done_only_scope(delivery_data) -> None:
    database, tenant, _, novel = delivery_data
    service = NovelExportService(database)
    options = await service.options(tenant_id=tenant.id, project_id=novel.id)
    assert options["volumes"][0]["selection"] == "partial"
    manifest = await service.export(
        tenant_id=tenant.id,
        project_id=novel.id,
        chapter_ids=("chapter-1",),
        form="working",
        idempotency_key="novel-export",
    )
    assert manifest.status == "succeeded"
    document = Document(io.BytesIO(manifest.artifact))
    assert "来信" in [paragraph.text for paragraph in document.paragraphs]
    with pytest.raises(NovelDeliveryError, match="completed"):
        await service.export(
            tenant_id=tenant.id,
            project_id=novel.id,
            chapter_ids=("chapter-2",),
            form="clean",
            idempotency_key="empty-chapter",
        )


@pytest.mark.asyncio
async def test_faithful_translation_is_export_only_and_preserves_block_contract(
    delivery_data,
) -> None:
    database, tenant, _, novel = delivery_data

    class Translator:
        async def translate(self, *, units, target_language, **kwargs):
            assert target_language == "en-US"
            unit = units[0]
            return (
                TranslationUnit(
                    titles={"volume_title": "Volume One", "chapter_title": "The Letter"},
                    blocks=tuple(
                        {**block, "text": f"translated: {block['text']}"} for block in unit.blocks
                    ),
                ),
            )

    service = NovelExportService(database, translator=Translator())
    manifest = await service.export(
        tenant_id=tenant.id,
        project_id=novel.id,
        chapter_ids=("chapter-1",),
        form="clean",
        translation_mode="faithful",
        target_language="en-US",
        idempotency_key="translated-export",
    )
    assert manifest.status == "succeeded"
    document = Document(io.BytesIO(manifest.artifact))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "The Letter" in paragraphs
    assert "translated: 雨落在窗沿，信纸发出轻响。" in paragraphs
    async with database.session() as session:
        original = (
            await session.scalars(
                select(NovelDocumentRevisionModel).where(
                    NovelDocumentRevisionModel.project_id == novel.id
                )
            )
        ).one()
        assert original.blocks[1]["text"] == "雨落在窗沿，信纸发出轻响。"
