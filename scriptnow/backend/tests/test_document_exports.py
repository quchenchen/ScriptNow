import io

import pytest
from docx import Document

from scriptnow.novel.contracts import NovelBlock
from scriptnow.novel.export import NovelExportChapter, render_novel_docx
from scriptnow.script.contracts import ScriptBlock
from scriptnow.script.export import ScriptExportScene, render_script_docx


def test_script_golden_document_has_distinct_china_and_hollywood_layouts() -> None:
    scene = ScriptExportScene(
        episode_title="第一集",
        scene_title="1. 雨夜来信",
        blocks=(
            ScriptBlock(para_id="p1", type="slugline", text="内景 客厅 夜"),
            ScriptBlock(para_id="p2", type="action", text="她拆开信封。"),
            ScriptBlock(para_id="p3", type="character", text="林夏"),
            ScriptBlock(para_id="p4", type="dialogue", text="这不可能。"),
            ScriptBlock(para_id="p5", type="transition", text="切至："),
        ),
    )
    china = Document(
        io.BytesIO(
            render_script_docx(project_name="金色项目", script_format="chinese", scenes=(scene,))
        )
    )
    hollywood = Document(
        io.BytesIO(
            render_script_docx(project_name="Golden", script_format="hollywood", scenes=(scene,))
        )
    )

    china_text = [p.text for p in china.paragraphs]
    hollywood_text = [p.text for p in hollywood.paragraphs]
    assert "《金色项目》" in china_text
    assert "第 1 集" in china_text
    assert "第 1-1 场 内景 客厅 夜" in china_text
    assert "人物：林夏" in china_text
    short = Document(
        io.BytesIO(
            render_script_docx(
                project_name="短剧项目",
                script_format="chinese-short",
                scenes=(scene,),
            )
        )
    )
    short_text = [p.text for p in short.paragraphs]
    assert "1-1 内景 客厅 夜" in short_text
    assert "出场人物：林夏" in short_text
    assert "▲她拆开信封。" in short_text
    assert "△她拆开信封。" in china_text
    assert "林夏：这不可能。" in china_text
    assert "Golden" in hollywood_text
    assert "INT. APARTMENT - NIGHT" not in hollywood_text
    assert "内景 客厅 夜" in hollywood_text
    assert "第 1 集" not in hollywood_text
    assert "人物：林夏" not in hollywood_text

    china_character = next(p for p in china.paragraphs if p.text == "林夏：这不可能。")
    hollywood_character = next(p for p in hollywood.paragraphs if p.text == "林夏")
    assert china_character.paragraph_format.left_indent.cm == pytest.approx(3.0, abs=0.01)
    assert hollywood_character.paragraph_format.left_indent.cm == pytest.approx(6.35, abs=0.01)
    assert hollywood.styles["Normal"].font.name == "Courier New"


def test_hollywood_export_keeps_spec_script_scene_heading_without_scene_numbers() -> None:
    scene = ScriptExportScene(
        episode_title="Pilot",
        scene_title="The call",
        episode_ordinal=1,
        scene_ordinal=7,
        blocks=(
            ScriptBlock(para_id="h1", type="slugline", text="INT. CLINIC - NIGHT"),
            ScriptBlock(para_id="h2", type="action", text="A phone rings in the empty ward."),
            ScriptBlock(para_id="h3", type="character", text="LIN"),
            ScriptBlock(para_id="h4", type="dialogue", text="I know what you did."),
        ),
    )
    document = Document(
        io.BytesIO(
            render_script_docx(project_name="Echo Clinic", script_format="hollywood", scenes=(scene,))
        )
    )
    text = [paragraph.text for paragraph in document.paragraphs]
    assert "INT. CLINIC - NIGHT" in text
    assert not any("1-7" in paragraph for paragraph in text)
    assert not any(paragraph.startswith("人物：") for paragraph in text)


def test_novel_golden_document_preserves_headings_quotes_and_page_breaks() -> None:
    chapters = tuple(
        NovelExportChapter(
            volume_title="第一卷",
            chapter_title=f"第 {index} 章",
            blocks=(
                NovelBlock(block_id=f"h{index}", type="heading", text="迟来的信"),
                NovelBlock(block_id=f"p{index}", type="prose", text="雨落在窗沿。"),
                NovelBlock(block_id=f"q{index}", type="quote", text="你必须记住。"),
            ),
        )
        for index in (1, 2)
    )
    document = Document(io.BytesIO(render_novel_docx(project_name="长夜", chapters=chapters)))

    assert [p.style.name for p in document.paragraphs if p.text == "迟来的信"] == [
        "Heading 3",
        "Heading 3",
    ]
    quote = next(p for p in document.paragraphs if p.text == "你必须记住。")
    assert quote.paragraph_format.left_indent.cm == pytest.approx(1.5, abs=0.01)
    xml = document._element.xml
    assert 'w:type="page"' in xml
