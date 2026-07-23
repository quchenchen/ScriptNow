import io
from dataclasses import dataclass
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from scriptflow_v7.script.contracts import ScriptBlock

ScriptFormat = Literal["chinese", "hollywood"]


@dataclass(frozen=True, slots=True)
class ScriptExportScene:
    episode_title: str
    scene_title: str
    blocks: tuple[ScriptBlock, ...]


def render_script_docx(
    *, project_name: str, script_format: ScriptFormat, scenes: tuple[ScriptExportScene, ...]
) -> bytes:
    """Render Script blocks without importing any Novel contract or formatter."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18 if script_format == "hollywood" else 2.54)
    section.right_margin = Cm(2.54)
    normal = document.styles["Normal"]
    normal.font.name = "Courier New" if script_format == "hollywood" else "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    title = document.add_paragraph(project_name)
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Courier New" if script_format == "hollywood" else "黑体"
    title.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.runs[0].font.size = Pt(22)

    current_episode = None
    for scene in scenes:
        if current_episode != scene.episode_title:
            if current_episode is not None:
                document.add_page_break()
            episode = document.add_paragraph(scene.episode_title)
            episode.style = document.styles["Heading 1"]
            episode.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_episode = scene.episode_title
        heading = document.add_paragraph(scene.scene_title)
        heading.style = document.styles["Heading 2"]
        for block in scene.blocks:
            paragraph = document.add_paragraph(block.text)
            paragraph.style = document.styles["Normal"]
            _format_script_paragraph(paragraph, block.type, script_format)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _format_script_paragraph(paragraph, block_type: str, script_format: ScriptFormat) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(0 if block_type in {"character", "dialogue"} else 6)
    fmt.line_spacing = 1.0 if script_format == "hollywood" else 1.25
    if block_type == "slugline":
        paragraph.runs[0].bold = True
        paragraph.runs[0].text = paragraph.runs[0].text.upper()
    elif block_type == "character":
        fmt.left_indent = Cm(6.35 if script_format == "hollywood" else 5.0)
        paragraph.runs[0].bold = True
    elif block_type == "dialogue":
        fmt.left_indent = Cm(3.8 if script_format == "hollywood" else 3.0)
        fmt.right_indent = Cm(3.8 if script_format == "hollywood" else 3.0)
    elif block_type == "transition":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.runs[0].bold = True
